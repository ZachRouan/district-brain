import io

import pytest

from corpus.chunking import chunk_text
from corpus.extractors import UnsupportedFormat, extract_text


def minimal_pdf(text: str) -> bytes:
    """Assemble a tiny but valid one-page PDF containing `text`."""
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R "
        b"/Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length %d >>\nstream\n" % len(stream) + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objects, 1):
        offsets.append(len(out))
        out += b"%d 0 obj\n" % i + obj + b"\nendobj\n"
    xref_pos = len(out)
    out += b"xref\n0 %d\n" % (len(objects) + 1) + b"0000000000 65535 f \n"
    for off in offsets:
        out += b"%010d 00000 n \n" % off
    out += b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF" % (len(objects) + 1, xref_pos)
    return bytes(out)


def test_txt_and_md(tmp_path):
    for name in ("policy.txt", "policy.md"):
        path = tmp_path / name
        path.write_text("Devices must remain silenced during class.")
        assert "silenced during class" in extract_text(path)


def test_html_strips_markup_and_scripts(tmp_path):
    path = tmp_path / "handbook.html"
    path.write_text(
        "<html><head><script>alert('x')</script><style>p{}</style></head>"
        "<body><h1>Dress code</h1><p>Hats are <b>not</b> permitted indoors.</p></body></html>"
    )
    text = extract_text(path)
    assert "Hats are not permitted indoors." in text.replace("\n", " ")
    assert "alert" not in text
    assert "<p>" not in text


def test_docx(tmp_path):
    import docx

    path = tmp_path / "minutes.docx"
    d = docx.Document()
    d.add_paragraph("The board approved the 2026-27 bell schedule.")
    d.save(path)
    assert "approved the 2026-27 bell schedule" in extract_text(path)


def test_pdf(tmp_path):
    path = tmp_path / "policy.pdf"
    path.write_bytes(minimal_pdf("Buses depart at 2:45 PM sharp."))
    assert "Buses depart at 2:45 PM sharp." in extract_text(path)


def test_unsupported_extension(tmp_path):
    path = tmp_path / "grades.xlsx"
    path.write_bytes(b"not really a spreadsheet")
    with pytest.raises(UnsupportedFormat):
        extract_text(path)


# --- Table enrichment: rows become sentence-shaped, context- and header-labeled,
#     each its own chunk, so tabular facts are retrievable by plain questions. ---

BELL_MD = (
    "# Bell Schedules 2026-27\n\n"
    "## Regular schedule\n\n"
    "| Period | Start | End |\n"
    "| --- | --- | --- |\n"
    "| First Period | 7:50 AM | 8:38 AM |\n"
    "| Second Period | 8:42 AM | 9:30 AM |\n"
)


def test_markdown_table_rows_are_enriched_with_heading_and_headers():
    text = extract_text(io.BytesIO(BELL_MD.encode()), filename="bell.md", title="Bell Schedules 2026-27")
    assert "Regular schedule: Period: First Period, Start: 7:50 AM, End: 8:38 AM" in text
    # Each row is its own chunk, not merged into a diluted table blob.
    chunks = chunk_text(text)
    assert "Regular schedule: Period: First Period, Start: 7:50 AM, End: 8:38 AM" in chunks
    assert "Regular schedule: Period: Second Period, Start: 8:42 AM, End: 9:30 AM" in chunks
    # The nearest heading is the context; the standalone title is not prepended
    # onto rows (it dilutes retrieval and already appears on every citation).
    assert "Bell Schedules 2026-27: Period:" not in text


def test_markdown_non_table_content_is_unchanged():
    md = "# Heading\n\nA plain paragraph about silenced devices.\n\nAnother paragraph."
    text = extract_text(io.BytesIO(md.encode()), filename="p.md")
    assert "A plain paragraph about silenced devices." in text
    assert "\f" not in text  # no table, no hard boundaries introduced


def test_html_table_with_headers_is_enriched():
    html = (
        "<h2>Regular schedule</h2>"
        "<table><tr><th>Period</th><th>Start</th><th>End</th></tr>"
        "<tr><td>First Period</td><td>7:50 AM</td><td>8:38 AM</td></tr></table>"
    )
    text = extract_text(io.BytesIO(html.encode()), filename="s.html", title="Bell Schedules")
    assert "Regular schedule: Period: First Period, Start: 7:50 AM, End: 8:38 AM" in text


def test_html_headerless_table_falls_back_to_pipes_with_context():
    html = "<h2>Room status</h2><table><tr><td>Gymnasium</td><td>Open</td></tr></table>"
    text = extract_text(io.BytesIO(html.encode()), filename="s.html")
    assert "Room status: Gymnasium | Open" in text


def test_docx_table_rows_are_enriched():
    import docx

    d = docx.Document()
    d.add_heading("Regular schedule", level=2)
    table = d.add_table(rows=3, cols=3)
    data = [
        ("Period", "Start", "End"),
        ("First Period", "7:50 AM", "8:38 AM"),
        ("Second Period", "8:42 AM", "9:30 AM"),
    ]
    for r, (a, b, c) in enumerate(data):
        table.rows[r].cells[0].text, table.rows[r].cells[1].text, table.rows[r].cells[2].text = a, b, c
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    text = extract_text(buf, filename="sched.docx", title="Bell Schedules")
    assert "Regular schedule: Period: First Period, Start: 7:50 AM, End: 8:38 AM" in text


def test_table_with_no_heading_falls_back_to_the_document_title():
    md = "| Period | Start |\n| --- | --- |\n| First Period | 7:50 AM |\n"
    text = extract_text(io.BytesIO(md.encode()), filename="bare.md", title="Bell Schedules 2026-27")
    assert "Bell Schedules 2026-27: Period: First Period, Start: 7:50 AM" in text


def test_html_text_in_bare_containers_is_not_dropped():
    """Word's "Save as HTML" and Google exports put body text in bare <div>/<span>
    elements with no <p>. That text must reach the corpus, not vanish silently."""
    html = (
        "<h1>Dress code</h1><div>Hats are not permitted indoors.</div>"
        "<span>Buses leave at 2:45.</span><div><p>Inside a paragraph.</p><div>Sibling div text.</div></div>"
        "<!-- a comment is not content -->"
    )
    text = extract_text(io.BytesIO(html.encode()), filename="export.html")
    assert "Hats are not permitted indoors." in text
    assert "Buses leave at 2:45." in text
    assert "Inside a paragraph." in text
    assert "Sibling div text." in text
    assert "comment" not in text


def test_html_loose_text_directly_in_body_is_kept():
    html = "<body>Loose opening line.<p>A paragraph.</p></body>"
    text = extract_text(io.BytesIO(html.encode()), filename="loose.html")
    assert "Loose opening line." in text
    assert "A paragraph." in text


def test_markdown_horizontal_rule_is_not_a_table():
    md = "Options: a | b\n\n---\n\nAfter the rule."
    text = extract_text(io.BytesIO(md.encode()), filename="rule.md")
    assert "Options: a | b" in text
    assert "After the rule." in text
    assert "\f" not in text
